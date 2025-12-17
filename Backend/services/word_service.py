from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any
from datetime import datetime
from google.cloud import storage
from config.settings import settings
import tempfile
import os

def create_word_document(analysis: Dict[str, Any], deal_id: str) -> str:
    """Create a Word document from analysis data and upload to GCS"""
    temp_path = None
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Investment Memo', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Company Overview Section
        company_overview = analysis.get('company_overview', {})
        doc.add_heading('Company Overview', 1)
        doc.add_paragraph(f"Company Name: {company_overview.get('name', 'N/A')}")
        doc.add_paragraph(f"Sector: {company_overview.get('sector', 'N/A')}")
        
        # Founders
        founders = company_overview.get('founders', [])
        if founders:
            doc.add_heading('Founders', 2)
            for founder in founders:
                if isinstance(founder, str):
                    doc.add_paragraph(f"Name: {founder}", style='List Bullet')
                else:
                    doc.add_paragraph(f"Name: {founder.get('name', 'N/A')}", style='List Bullet')
                    doc.add_paragraph(f"Education: {founder.get('education', 'N/A')}", style='List Bullet 2')
                    doc.add_paragraph(f"Background: {founder.get('professional_background', 'N/A')}", style='List Bullet 2')
                    if founder.get('previous_ventures'):
                        doc.add_paragraph(f"Previous Ventures: {founder.get('previous_ventures')}", style='List Bullet 2')
        
        # Technologies
        if company_overview.get('technologies_used'):
            doc.add_heading('Technologies Used', 2)
            doc.add_paragraph(company_overview.get('technologies_used', 'N/A'))
        
        # Key Problems Solved
        key_problems = company_overview.get('key_problems_solved', [])
        if key_problems:
            doc.add_heading('Key Problems Solved', 2)
            for problem in key_problems:
                doc.add_paragraph(problem, style='List Bullet')
        
        # Market Analysis Section
        market_analysis = analysis.get('market_analysis', {})
        doc.add_heading('Market Analysis', 1)
        
        # Industry Size and Growth
        industry = market_analysis.get('industry_size_and_growth', {})
        if industry:
            doc.add_heading('Industry Size and Growth', 2)
            
            tam = industry.get('total_addressable_market', {})
            if tam:
                doc.add_paragraph('Total Addressable Market (TAM):', style='List Bullet')
                doc.add_paragraph(f"Value: {tam.get('value', 'N/A')}", style='List Bullet 2')
                doc.add_paragraph(f"CAGR: {tam.get('cagr', 'N/A')}", style='List Bullet 2')
                doc.add_paragraph(f"Source: {tam.get('source', 'N/A')}", style='List Bullet 2')
            
            som = industry.get('serviceable_obtainable_market', {})
            if som:
                doc.add_paragraph('Serviceable Obtainable Market (SOM):', style='List Bullet')
                doc.add_paragraph(f"Value: {som.get('value', 'N/A')}", style='List Bullet 2')
                doc.add_paragraph(f"CAGR: {som.get('cagr', 'N/A')}", style='List Bullet 2')
                doc.add_paragraph(f"Projection: {som.get('projection', 'N/A')}", style='List Bullet 2')
            
            if industry.get('commentary'):
                doc.add_paragraph(f"Commentary: {industry.get('commentary')}")
        
        # Sub-segment Opportunities
        sub_segments = market_analysis.get('sub_segment_opportunities', [])
        if sub_segments:
            doc.add_heading('Sub-segment Opportunities', 2)
            for segment in sub_segments:
                doc.add_paragraph(segment, style='List Bullet')
        
        # Competitors
        competitors = market_analysis.get('competitor_details', [])
        if competitors:
            doc.add_heading('Competitor Analysis', 2)
            for comp in competitors:
                doc.add_paragraph(f"• {comp.get('name', 'N/A')}", style='Heading 3')
                if comp.get('headquarters'):
                    doc.add_paragraph(f"Headquarters: {comp.get('headquarters')}", style='List Bullet')
                if comp.get('founding_year'):
                    doc.add_paragraph(f"Founded: {comp.get('founding_year')}", style='List Bullet')
                if comp.get('total_funding_raised'):
                    doc.add_paragraph(f"Funding: {comp.get('total_funding_raised')}", style='List Bullet')
                if comp.get('business_model'):
                    doc.add_paragraph(f"Business Model: {comp.get('business_model')}", style='List Bullet')
                if comp.get('current_arr'):
                    doc.add_paragraph(f"ARR: {comp.get('current_arr')}", style='List Bullet')
        
        # Industry Reports
        reports = market_analysis.get('reports', [])
        if reports:
            doc.add_heading('Industry Reports', 2)
            for report in reports:
                doc.add_paragraph(f"• {report.get('title', 'N/A')}", style='List Bullet')
                doc.add_paragraph(f"Source: {report.get('source_name', 'N/A')}", style='List Bullet 2')
                if report.get('summary'):
                    doc.add_paragraph(f"Summary: {report.get('summary')}", style='List Bullet 2')
        
        # Business Model Section (LIST)
        business_models = analysis.get('business_model', [])
        if business_models:
            doc.add_heading('Business Model', 1)
            for idx, bm in enumerate(business_models, 1):
                doc.add_heading(f"Revenue Stream {idx}: {bm.get('revenue_streams', 'N/A')}", 2)
                
                if bm.get('description'):
                    doc.add_paragraph(f"Description: {bm.get('description')}")
                
                if bm.get('target_audience'):
                    doc.add_paragraph(f"Target Audience: {bm.get('target_audience')}")
                
                if bm.get('percentage_contribution'):
                    doc.add_paragraph(f"Revenue Contribution: {bm.get('percentage_contribution')}")
                
                if bm.get('pricing'):
                    doc.add_paragraph(f"Pricing: {bm.get('pricing')}")
                
                # Unit Economics
                unit_econ = bm.get('unit_economics', {})
                if unit_econ:
                    doc.add_paragraph('Unit Economics:', style='List Bullet')
                    doc.add_paragraph(f"CAC: {unit_econ.get('customer_acquisition_cost_CAC', 'N/A')}", style='List Bullet 2')
                    doc.add_paragraph(f"LTV: {unit_econ.get('lifetime_value_LTV', 'N/A')}", style='List Bullet 2')
                    doc.add_paragraph(f"LTV/CAC Ratio: {unit_econ.get('LTV_CAC_Ratio', 'N/A')}", style='List Bullet 2')
                
                if bm.get('scalability'):
                    doc.add_paragraph(f"Scalability: {bm.get('scalability')}")
                
                additional_opps = bm.get('additional_revenue_opportunities', [])
                if additional_opps:
                    doc.add_paragraph('Additional Revenue Opportunities:', style='List Bullet')
                    for opp in additional_opps:
                        doc.add_paragraph(opp, style='List Bullet 2')
        
        # Financials Section
        financials = analysis.get('financials', {})
        doc.add_heading('Financials', 1)
        
        # ARR/MRR
        arr_mrr = financials.get('arr_mrr', {})
        if arr_mrr:
            doc.add_heading('ARR/MRR', 2)
            doc.add_paragraph(f"Current ARR: {arr_mrr.get('current_booked_arr', 'N/A')}")
            doc.add_paragraph(f"Current MRR: {arr_mrr.get('current_mrr', 'N/A')}")
        
        # Burn and Runway
        burn_runway = financials.get('burn_and_runway', {})
        if burn_runway:
            doc.add_heading('Burn and Runway', 2)
            doc.add_paragraph(f"Funding Ask: {burn_runway.get('funding_ask', 'N/A')}")
            doc.add_paragraph(f"Stated Runway: {burn_runway.get('stated_runway', 'N/A')}")
            doc.add_paragraph(f"Net Burn: {burn_runway.get('implied_net_burn', 'N/A')}")
            doc.add_paragraph(f"Gross Margin: {burn_runway.get('gross_margin', 'N/A')}")
            doc.add_paragraph(f"CM1: {burn_runway.get('cm1', 'N/A')}")
            doc.add_paragraph(f"CM2: {burn_runway.get('cm2', 'N/A')}")
            doc.add_paragraph(f"CM3: {burn_runway.get('cm3', 'N/A')}")
        
        # Funding History
        if financials.get('funding_history'):
            doc.add_heading('Funding History', 2)
            doc.add_paragraph(financials.get('funding_history'))
        
        # Valuation
        if financials.get('valuation_rationale'):
            doc.add_heading('Valuation Rationale', 2)
            doc.add_paragraph(financials.get('valuation_rationale'))
        
        # Projections
        projections = financials.get('projections', [])
        if projections:
            doc.add_heading('Revenue Projections', 2)
            for proj in projections:
                doc.add_paragraph(f"{proj.get('year', 'N/A')}: {proj.get('revenue', 'N/A')}", style='List Bullet')
        
        # Claims Analysis Section
        claims = analysis.get('claims_analysis', [])
        if claims:
            doc.add_heading('Claims Analysis', 1)
            for idx, claim in enumerate(claims, 1):
                doc.add_heading(f"Claim {idx}", 2)
                doc.add_paragraph(f"Claim: {claim.get('claim', 'N/A')}")
                doc.add_paragraph(f"Analysis Method: {claim.get('analysis_method', 'N/A')}")
                doc.add_paragraph(f"Simulated Probability: {claim.get('simulated_probability', 'N/A')}")
                doc.add_paragraph(f"Result: {claim.get('result', 'N/A')}")
        
        # Risk Metrics Section
        risk_metrics = analysis.get('risk_metrics', {})
        if risk_metrics:
            doc.add_heading('Risk Assessment', 1)
            doc.add_paragraph(f"Composite Risk Score: {risk_metrics.get('composite_risk_score', 'N/A')}")
            doc.add_paragraph(f"Interpretation: {risk_metrics.get('score_interpretation', 'N/A')}")
            doc.add_paragraph(f"Justification: {risk_metrics.get('narrative_justification', 'N/A')}")
        
        # Risks and Mitigation Section
        risks = analysis.get('risks_and_mitigation', [])
        if risks:
            doc.add_heading('Risks and Mitigation Strategies', 1)
            for risk in risks:
                doc.add_heading(risk.get('risk', 'Risk'), 2)
                if risk.get('description'):
                    doc.add_paragraph(f"Description: {risk.get('description')}")
                if risk.get('likelihood'):
                    doc.add_paragraph(f"Likelihood: {risk.get('likelihood')}")
                if risk.get('impact'):
                    doc.add_paragraph(f"Impact: {risk.get('impact')}")
                if risk.get('mitigation'):
                    doc.add_paragraph(f"Mitigation: {risk.get('mitigation')}")
        
        # Conclusion Section
        conclusion = analysis.get('conclusion', {})
        if conclusion:
            doc.add_heading('Conclusion', 1)
            doc.add_paragraph(conclusion.get('overall_attractiveness', 'N/A'))
        
        # Add footer with generation date
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC"
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Create temporary file using tempfile module (cross-platform)
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as tmp_file:
            temp_path = tmp_file.name
            doc.save(temp_path)
        
        print(f"Word document saved to temporary file: {temp_path}")
        
        # Upload to GCS
        storage_client = storage.Client(project=settings.GCP_PROJECT_ID)
        bucket = storage_client.bucket(settings.GCS_BUCKET_NAME)
        blob_path = f"deals/{deal_id}/memo.docx"
        blob = bucket.blob(blob_path)
        blob.upload_from_filename(temp_path)
        
        gcs_url = f"gs://{settings.GCS_BUCKET_NAME}/{blob_path}"
        print(f"Word document uploaded to: {gcs_url}")
        
        return gcs_url
    
    except Exception as e:
        print(f"Error creating Word document: {str(e)}")
        raise Exception(f"Failed to create document: {str(e)}")
    
    finally:
        # Clean up temporary file
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
                print(f"Temporary file cleaned up: {temp_path}")
            except Exception as e:
                print(f"Warning: Could not delete temporary file {temp_path}: {str(e)}")
